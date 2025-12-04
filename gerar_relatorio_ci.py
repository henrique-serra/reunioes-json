#!/usr/bin/env python3
"""
Script para gerar relatório anual da Comissão de Serviços de Infraestrutura (CI)
Lê dados da planilha consolidada e gera documento Word formatado
"""

import sys
import os
import subprocess
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl


def definir_bordas_tabela(table):
    """Define bordas para todas as células da tabela"""
    tbl = table._element
    tblPr = tbl.tblPr

    # Criar elemento de bordas se não existir
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    # Definir cada tipo de borda
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Tamanho da borda
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)


def mesclar_celulas(cell1, cell2):
    """Mescla duas células horizontalmente"""
    cell1.merge(cell2)


def criar_tabela_formatada(doc, dados_tabela, titulo_tabela=None):
    """
    Cria uma tabela formatada no documento Word

    Args:
        doc: Documento Word
        dados_tabela: Lista de tuplas (coluna1, coluna2)
        titulo_tabela: Título opcional da tabela
    """
    # Adicionar parágrafo vazio antes da tabela
    doc.add_paragraph()

    # Criar tabela (adicionar 1 linha extra se houver título)
    num_rows = len(dados_tabela) + (1 if titulo_tabela else 0)
    table = doc.add_table(rows=num_rows, cols=2)
    table.allow_autofit = False

    # Definir largura das colunas
    table.columns[0].width = Inches(3.5)
    table.columns[1].width = Inches(2.0)

    row_offset = 0

    # Adicionar linha de título se existir
    if titulo_tabela:
        row = table.rows[0]
        # Mesclar as duas células
        mesclar_celulas(row.cells[0], row.cells[1])

        # Adicionar texto do título
        cell = row.cells[0]
        cell.text = titulo_tabela

        # Formatação do título: negrito e centralizado
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Sombreamento cinza claro
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9D9D9')
        cell._element.get_or_add_tcPr().append(shading_elm)

        row_offset = 1

    # Preencher dados
    for idx, (col1, col2) in enumerate(dados_tabela):
        row = table.rows[idx + row_offset]

        # Célula da primeira coluna
        cell1 = row.cells[0]
        cell1.text = str(col1) if col1 is not None else ''

        # Célula da segunda coluna
        cell2 = row.cells[1]
        if col2 is not None:
            cell2.text = str(col2) if not isinstance(col2, (int, float)) else str(int(col2))
        else:
            cell2.text = ''

        # Formatação do cabeçalho de dados (primeira linha de dados)
        if idx == 0:
            for cell in row.cells:
                # Negrito e centralizado
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Sombreamento cinza claro
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'D9D9D9')
                cell._element.get_or_add_tcPr().append(shading_elm)
        else:
            # Formatação das demais linhas - TODAS alinhadas à esquerda
            cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in cell1.paragraphs[0].runs:
                run.font.size = Pt(11)

            cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in cell2.paragraphs[0].runs:
                run.font.size = Pt(11)

    # Aplicar bordas
    definir_bordas_tabela(table)

    return table


def executar_consolidacao(ano, dir_base):
    """
    Executa o script consolidar_votacoes_ci.py para gerar a planilha consolidada

    Args:
        ano: Ano a consolidar
        dir_base: Diretório base onde está o script

    Returns:
        True se executado com sucesso, False caso contrário
    """
    script_consolidacao = os.path.join(dir_base, 'consolidar_votacoes_ci.py')

    if not os.path.exists(script_consolidacao):
        print(f"ERRO: Script de consolidação não encontrado: {script_consolidacao}")
        return False

    print(f"\n{'='*70}")
    print(f"EXECUTANDO CONSOLIDAÇÃO DE DADOS PARA O ANO {ano}")
    print(f"{'='*70}\n")

    try:
        # Executar o script de consolidação passando o ano via stdin
        processo = subprocess.Popen(
            [sys.executable, script_consolidacao],
            stdin=subprocess.PIPE,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            cwd=dir_base
        )

        # Enviar o ano para o stdin do processo
        # Também enviar "s" para responder automaticamente "sim" ao download
        stdout, stderr = processo.communicate(input=f"{ano}\ns\n")

        # Mostrar a saída do script
        if stdout:
            print(stdout)

        if stderr:
            print("ERROS/AVISOS:", file=sys.stderr)
            print(stderr, file=sys.stderr)

        if processo.returncode == 0:
            print(f"\n{'='*70}")
            print("CONSOLIDAÇÃO CONCLUÍDA COM SUCESSO")
            print(f"{'='*70}\n")
            return True
        else:
            print(f"\n{'='*70}")
            print(f"ERRO: Consolidação falhou com código de retorno {processo.returncode}")
            print(f"{'='*70}\n")
            return False

    except Exception as e:
        print(f"ERRO ao executar script de consolidação: {e}")
        return False


def extrair_tabelas_excel(caminho_excel):
    """
    Extrai as tabelas da aba 'Resumo Anual' do arquivo Excel

    Returns:
        Lista de dicionários com 'titulo' e 'dados' de cada tabela
    """
    wb = openpyxl.load_workbook(caminho_excel, data_only=True)

    if 'Resumo Anual' not in wb.sheetnames:
        raise ValueError("Aba 'Resumo Anual' não encontrada no arquivo Excel")

    ws = wb['Resumo Anual']

    tabelas = []
    tabela_atual = None

    for row in ws.iter_rows(values_only=True):
        # Verificar se é uma linha vazia
        if all(cell is None or str(cell).strip() == '' for cell in row):
            if tabela_atual and tabela_atual['dados']:
                tabelas.append(tabela_atual)
                tabela_atual = None
            continue

        # Verificar se é um título de tabela (primeira célula preenchida, segunda vazia)
        if row[0] and (not row[1] or str(row[1]).strip() == ''):
            if tabela_atual and tabela_atual['dados']:
                tabelas.append(tabela_atual)
            tabela_atual = {
                'titulo': str(row[0]).strip(),
                'dados': []
            }
        # Se já existe uma tabela iniciada, adicionar os dados
        elif tabela_atual:
            tabela_atual['dados'].append((row[0], row[1]))

    # Adicionar última tabela se existir
    if tabela_atual and tabela_atual['dados']:
        tabelas.append(tabela_atual)

    wb.close()
    return tabelas


def gerar_relatorio(ano):
    """
    Gera o relatório anual da CI

    Args:
        ano: Ano do relatório (string ou int)
    """
    ano = str(ano)

    # Caminhos dos arquivos
    dir_base = os.path.dirname(os.path.abspath(__file__))
    arquivo_excel = os.path.join(dir_base, f'votacoes_ci_consolidadas_{ano}.xlsx')
    arquivo_template = os.path.join(dir_base, 'relatorio_ci_ano.docx')
    arquivo_saida = os.path.join(dir_base, f'relatorio_ci_{ano}.docx')

    # Verificar se a planilha consolidada existe
    if not os.path.exists(arquivo_excel):
        print(f"\nArquivo de dados não encontrado: {arquivo_excel}")
        print(f"Será necessário executar a consolidação de dados primeiro.\n")

        # Executar script de consolidação
        sucesso = executar_consolidacao(ano, dir_base)

        if not sucesso:
            print(f"ERRO: Não foi possível consolidar os dados.")
            sys.exit(1)

        # Verificar novamente se o arquivo foi criado
        if not os.path.exists(arquivo_excel):
            print(f"ERRO: Arquivo consolidado não foi criado: {arquivo_excel}")
            sys.exit(1)

    if not os.path.exists(arquivo_template):
        print(f"ERRO: Template não encontrado: {arquivo_template}")
        sys.exit(1)

    print(f"Gerando relatório para o ano {ano}...")
    print(f"Lendo dados de: {arquivo_excel}")

    # Extrair tabelas do Excel
    try:
        tabelas = extrair_tabelas_excel(arquivo_excel)
        print(f"✓ {len(tabelas)} tabelas extraídas da planilha")
    except Exception as e:
        print(f"ERRO ao ler planilha: {e}")
        sys.exit(1)

    # Criar documento a partir do template
    doc = Document(arquivo_template)

    # Substituir placeholders no texto
    data_atual = datetime.now().strftime('%d/%m/%Y')

    for paragraph in doc.paragraphs:
        if '{ano}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{ano}', ano)
        if '{data atual no formato DD/MM/YYYY}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{data atual no formato DD/MM/YYYY}', data_atual)

        # Encontrar parágrafo onde inserir as tabelas
        if '{ inserir tabelas aqui }' in paragraph.text:
            # Limpar o texto do placeholder
            paragraph.text = ''

            # Inserir todas as tabelas após este parágrafo
            # Precisamos inserir no local correto do documento
            p_element = paragraph._element
            p_parent = p_element.getparent()
            p_index = p_parent.index(p_element)

            # Adicionar cada tabela (excluindo a última se for a tabela vazia de Emendas)
            tabelas_filtradas = []
            for tabela_info in tabelas:
                # Filtrar tabela vazia de "EMENDAS AO ORÇAMENTO"
                if tabela_info['titulo'] == 'EMENDAS AO ORÇAMENTO' and len(tabela_info['dados']) == 1:
                    continue
                tabelas_filtradas.append(tabela_info)

            # Inserir tabelas filtradas
            for idx, tabela_info in enumerate(tabelas_filtradas):
                # Criar parágrafo vazio antes da tabela (espaçamento)
                if idx > 0:
                    novo_p = doc.add_paragraph()
                    p_parent.insert(p_index + 1, novo_p._element)
                    p_index += 1

                # Criar a tabela
                table = criar_tabela_formatada(doc, tabela_info['dados'], tabela_info['titulo'])

                # Inserir tabela no local correto
                p_parent.insert(p_index + 1, table._element)
                p_index += 1

            # Adicionar tabela de EMENDAS AO ORÇAMENTO manualmente
            novo_p = doc.add_paragraph()
            p_parent.insert(p_index + 1, novo_p._element)
            p_index += 1

            dados_emendas = [
                ('Tipo de matéria', 'Quantidade'),
                ('PLOA', ''),
                ('PLDO', '')
            ]
            table_emendas = criar_tabela_formatada(doc, dados_emendas, 'EMENDAS AO ORÇAMENTO')
            p_parent.insert(p_index + 1, table_emendas._element)
            p_index += 1

    # Salvar documento
    doc.save(arquivo_saida)
    print(f"✓ Relatório gerado com sucesso: {arquivo_saida}")


def main():
    """Função principal"""
    print("=" * 60)
    print("Gerador de Relatório Anual - Comissão de Infraestrutura")
    print("=" * 60)
    print()

    # Solicitar o ano ao usuário
    ano = input("Digite o ano do relatório que deseja gerar: ").strip()

    # Validar ano
    try:
        ano_int = int(ano)
        if ano_int < 2000 or ano_int > 2100:
            print("ERRO: Ano inválido. Digite um ano entre 2000 e 2100.")
            sys.exit(1)
    except ValueError:
        print("ERRO: Digite um ano válido (números apenas).")
        sys.exit(1)

    print()

    # Gerar relatório
    gerar_relatorio(ano)

    print()
    print("=" * 60)
    print("Processo concluído!")
    print("=" * 60)


if __name__ == "__main__":
    main()
